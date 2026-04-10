"""Document processing service."""

import asyncio
import hashlib
import json
import shutil
from pathlib import Path
from typing import Optional

from config import settings
from models.document_store import DocumentStore
from models.schemas import DocumentStatus
from lib.pageindex_wrapper import PageIndexWrapper

# Import LightRAG wrapper
try:
    from lib.lightrag import LightRAGWrapper

    LIGHTRAG_AVAILABLE = True
except ImportError:
    LIGHTRAG_AVAILABLE = False
    LightRAGWrapper = None

# Import HiRAG wrapper
try:
    from lib.hirag_wrapper import HiRAGWrapper

    HIRAG_AVAILABLE = True
except ImportError:
    HIRAG_AVAILABLE = False
    HiRAGWrapper = None

# Import Hybrid Search wrapper
try:
    from lib.hybrid_search import HybridSearchWrapper

    HYBRID_SEARCH_AVAILABLE = True
except ImportError:
    HYBRID_SEARCH_AVAILABLE = False
    HybridSearchWrapper = None


def calculate_content_hash(content: bytes) -> str:
    r"""Calculate SHA-256 hash of normalized content.

    Normalizes line endings to ensure cross-platform consistency:
    - \r\n (Windows) → \n
    - \r (old Mac) → \n
    - \n (Unix) stays \n

    Args:
        content: Raw file content as bytes

    Returns:
        64-character hexadecimal SHA-256 hash string
    """
    # Normalize line endings for cross-platform consistency
    normalized = content.replace(b"\r\n", b"\n").replace(b"\r", b"\n")
    return hashlib.sha256(normalized).hexdigest()


class DocumentService:
    """文档处理服务 — 并行构建多策略索引.

    文档上传后并行执行 4 种分块/索引策略（对应 Weaviate 文章的策略编号）：
    - PageIndex: S3 + S5 + S8 (文档结构 + LLM + 层次化分块)
    - LightRAG: S4 (语义分块，知识图谱驱动)
    - HiRAG: S8 (层次化分块，GMM 聚类)
    - Hybrid Search: S1 + S2 (固定大小 + 递归分块)

    策略详情：docs/chunking_strategy_mapping.md
    """

    def __init__(
        self,
        store: DocumentStore,
        pageindex_wrapper: PageIndexWrapper | None = None,
        lightrag_wrapper: Optional["LightRAGWrapper"] = None,
        hirag_wrapper: Optional["HiRAGWrapper"] = None,
        hybrid_search_wrapper: Optional["HybridSearchWrapper"] = None,
    ):
        """Initialize document service."""
        self.store = store
        self.pageindex = pageindex_wrapper or PageIndexWrapper()
        self.lightrag = lightrag_wrapper
        self.hirag = hirag_wrapper
        self.hybrid_search = hybrid_search_wrapper

    async def process_document(
        self,
        doc_id: str,
        file_path: str,
        file_format: str,
        content_hash: str | None = None,
        file_size: int | None = None,
    ) -> tuple[bool, list[str], dict[str, str]]:
        """Process document and build tree index - 并行构建多策略索引.

        Args:
            doc_id: 文档 ID
            file_path: 上传文件路径
            file_format: 文件格式
            content_hash: 内容哈希（可选）
            file_size: 文件大小（可选）

        Returns:
            Tuple of (success, available_indexes, failed_indexes)
        """
        available_indexes: list[str] = []
        failed_indexes: dict[str, str] = {}

        try:
            await self.store.update_status(doc_id, DocumentStatus.PROCESSING)

            storage_dir = settings.storage_path / doc_id
            storage_dir.mkdir(parents=True, exist_ok=True)

            # Copy original file
            original_path = storage_dir / f"original.{file_format}"
            shutil.copy(file_path, original_path)

            # Extract original filename from temp path "temp_{doc_id}_{filename}"
            temp_name = Path(file_path).name
            original_filename = temp_name[len(f"temp_{doc_id}_"):]

            # Delete temp upload file now that it has been copied
            temp = Path(file_path)
            if temp.exists() and temp.name.startswith("temp_"):
                temp.unlink()

            # 提取文本（根据格式）
            text = await self._extract_text(original_path, file_format)

            # 并行构建索引 - 带个别错误处理
            index_builders = []

            if self.pageindex:
                index_builders.append(
                    ("pageindex", self._build_pageindex(doc_id, original_path, file_format, storage_dir, original_filename))
                )

            if self.lightrag and LIGHTRAG_AVAILABLE:
                index_builders.append(("lightrag", self._build_lightrag(doc_id, text)))

            if self.hirag and HIRAG_AVAILABLE:
                index_builders.append(("hirag", self._build_hirag(doc_id, text)))

            if self.hybrid_search and HYBRID_SEARCH_AVAILABLE:
                index_builders.append(("hybrid_search", self._build_hybrid_search(doc_id, text, storage_dir)))

            # 并行执行所有索引构建，跟踪成功/失败
            if index_builders:
                names = [name for name, _ in index_builders]
                coros = [coro for _, coro in index_builders]
                results = await asyncio.gather(*coros, return_exceptions=True)
                for index_name, result in zip(names, results):
                    if isinstance(result, Exception):
                        error_msg = str(result)
                        failed_indexes[index_name] = error_msg
                        print(f"[DocumentService] {index_name} index failed for {doc_id}: {error_msg}")
                    else:
                        available_indexes.append(index_name)
                        print(f"[DocumentService] {index_name} index built successfully for {doc_id}")

            # 更新文档索引状态
            await self.store.update_indexes(doc_id, available_indexes, failed_indexes)

            # 只要有至少一个索引成功，就标记为 COMPLETED
            if available_indexes:
                await self.store.update_status(doc_id, DocumentStatus.COMPLETED)
                success = True
            else:
                await self.store.update_status(doc_id, DocumentStatus.FAILED, "All indexes failed")
                success = False

            return success, available_indexes, failed_indexes

        except Exception as e:
            error_msg = str(e)
            await self.store.update_status(doc_id, DocumentStatus.FAILED, error_msg)
            failed_indexes["_process"] = error_msg
            return False, available_indexes, failed_indexes

    async def _extract_text(self, file_path: Path, file_format: str) -> str:
        """从文件中提取纯文本."""
        if file_format == "txt" or file_format == "md":
            with open(file_path, encoding="utf-8") as f:
                return f.read()
        elif file_format == "docx":
            # 复用 _convert_docx_to_markdown 的逻辑
            return self._convert_docx_to_markdown(file_path)
        elif file_format == "pdf":
            # 使用 PyMuPDF 提取文本
            try:
                import fitz  # PyMuPDF

                doc = fitz.open(file_path)
                text_parts = []
                for page in doc:
                    text_parts.append(page.get_text())
                doc.close()
                return "\n\n".join(text_parts)
            except ImportError as err:
                raise ImportError("PyMuPDF is required for PDF text extraction") from err
        else:
            raise ValueError(f"Unsupported format for text extraction: {file_format}")

    async def _build_pageindex(
        self, doc_id: str, original_path: Path, file_format: str, storage_dir: Path,
        original_filename: str | None = None,
    ) -> dict:
        """构建 PageIndex 索引."""
        print(f"[DocumentService] Building PageIndex for {doc_id}")

        # 根据格式路由到 PageIndex 不同函数
        if file_format == "pdf":
            tree = await self.pageindex.build_tree_from_pdf(str(original_path), storage_dir)
        elif file_format == "docx":
            # 将 DOCX 转换为 Markdown
            md_content = self._convert_docx_to_markdown(original_path)
            temp_md = storage_dir / "temp.md"
            with open(temp_md, "w", encoding="utf-8") as f:
                f.write(md_content)
            tree = await self.pageindex.build_tree_from_markdown(str(temp_md), storage_dir)
            temp_md.unlink()  # 删除临时文件
        elif file_format in ["md", "txt"]:
            # 对于 txt 文件，先转换为临时 markdown 文件
            # 因为 md_to_tree 需要 .md 扩展名
            if file_format == "txt":
                temp_md = storage_dir / "temp.md"
                shutil.copy(original_path, temp_md)
                tree = await self.pageindex.build_tree_from_markdown(str(temp_md), storage_dir)
                temp_md.unlink()  # 删除临时文件
            else:
                tree = await self.pageindex.build_tree_from_markdown(
                    str(original_path), storage_dir
                )
        else:
            raise ValueError(f"Unsupported format: {file_format}")

        # Fix doc_name to use the real uploaded filename instead of "original.<ext>"
        if original_filename:
            tree["doc_name"] = original_filename

        # Save tree to JSON
        tree_path = storage_dir / "tree.json"
        with open(tree_path, "w", encoding="utf-8") as f:
            json.dump(tree, f, indent=2, ensure_ascii=False)

        print(f"[DocumentService] PageIndex built for {doc_id}")
        return tree

    async def _build_lightrag(self, document_id: str, text: str) -> dict:
        """构建 LightRAG 索引."""
        print(f"[DocumentService] Building LightRAG index for {document_id}")

        # 确保 LightRAG 已初始化
        await self.lightrag.initialize()

        # 调用包装器（内部使用 lightrag-hku 的 ainsert）
        result = await self.lightrag.index_document(
            document_id=document_id,
            text=text,
        )

        print(f"[DocumentService] LightRAG index built for {document_id}")
        return result

    async def _build_hirag(self, document_id: str, text: str) -> dict:
        """构建 HiRAG 索引."""
        print(f"[DocumentService] Building HiRAG index for {document_id}")

        # 确保 HiRAG 已初始化
        await self.hirag.initialize()

        # 调用包装器（内部使用 HiRAG 的 insert）
        result = await self.hirag.index_document(
            document_id=document_id,
            text=text,
        )

        print(f"[DocumentService] HiRAG index built for {document_id}")
        return result

    async def _build_hybrid_search(
        self, document_id: str, text: str, storage_dir: Path
    ) -> dict:
        """构建 Hybrid Search 索引."""
        print(f"[DocumentService] Building Hybrid Search index for {document_id}")

        # 确保 Hybrid Search 已初始化
        await self.hybrid_search.initialize()

        # 获取原始文件名作为 metadata
        original_path = storage_dir / "original"
        filename = "unknown"
        for ext in [".pdf", ".docx", ".md", ".txt"]:
            candidate = original_path.with_suffix(ext)
            if candidate.exists():
                filename = candidate.name
                break

        result = await self.hybrid_search.index_document(
            document_id=document_id,
            text=text,
            metadata={"filename": filename},
        )

        print(
            f"[DocumentService] Hybrid Search index built for {document_id}: "
            f"{result['chunks_count']} chunks"
        )
        return result

    def get_tree(self, doc_id: str) -> dict | None:
        """Load tree structure from storage."""
        tree_path = settings.storage_path / doc_id / "tree.json"
        if tree_path.exists():
            with open(tree_path, encoding="utf-8") as f:
                return json.load(f)
        return None

    def _convert_docx_to_markdown(self, docx_path: Path) -> str:
        """将 DOCX 文件转换为 Markdown 格式.

        使用 python-docx 提取文档内容并转换为简单的 Markdown。
        对于大表格，会尝试按内容分组以提高可搜索性。
        """
        try:
            from docx import Document
        except ImportError as err:
            raise ImportError(
                "python-docx is required for DOCX support. Install it with: pip install python-docx"
            ) from err

        doc = Document(docx_path)
        markdown_lines = []

        # 找到第一个非空段落文本，作为文档标题（仅当它不是已有标题样式时）
        injected_title_text: str | None = None
        for paragraph in doc.paragraphs[:5]:
            text = paragraph.text.strip()
            if text:
                style_name = paragraph.style.name.lower()
                # 如果第一段已经是标题样式，不需要额外注入
                if not any(f"heading {i}" in style_name for i in range(1, 7)):
                    markdown_lines.append(f"# {text}\n\n")
                    injected_title_text = text
                break

        # 处理段落
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            # 根据样式判断标题级别
            style_name = paragraph.style.name.lower()
            if "heading 1" in style_name:
                markdown_lines.append(f"# {text}\n\n")
            elif "heading 2" in style_name:
                markdown_lines.append(f"## {text}\n\n")
            elif "heading 3" in style_name:
                markdown_lines.append(f"### {text}\n\n")
            elif "heading 4" in style_name:
                markdown_lines.append(f"#### {text}\n\n")
            elif "heading 5" in style_name:
                markdown_lines.append(f"##### {text}\n\n")
            elif "heading 6" in style_name:
                markdown_lines.append(f"###### {text}\n\n")
            else:
                # 跳过已经手动注入为标题的段落（只跳过一次）
                if injected_title_text is not None and text == injected_title_text:
                    injected_title_text = None
                    continue
                markdown_lines.append(f"{text}\n\n")

        # 处理表格
        for table_idx, table in enumerate(doc.tables):
            # 为每个表格添加一个标题（如果文档中没有明确的表格标题）
            if len(doc.tables) > 1:
                markdown_lines.append(f"## 表格 {table_idx + 1}\n\n")

            # 尝试智能分组大表格（如果表格有"区域"或类似的分组列）
            if len(table.rows) > 10:
                grouped_content = self._try_group_table(table)
                if grouped_content:
                    markdown_lines.extend(grouped_content)
                    continue

            # 标准表格转换
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
                markdown_lines.append("| " + " | ".join(cells) + " |\n")
                # 添加表格分隔线
                if i == 0:
                    markdown_lines.append("| " + " | ".join(["---"] * len(cells)) + " |\n")
            markdown_lines.append("\n")

        return "".join(markdown_lines)

    def _try_group_table(self, table) -> list:
        """尝试按某一列的值对表格进行分组（例如按区域分组）.

        返回分组后的 Markdown 行列表，如果无法分组则返回 None.
        """
        if not table.rows or len(table.rows) < 2:
            return None

        # 获取表头
        header_row = table.rows[0]
        headers = [cell.text.strip() for cell in header_row.cells]

        # 查找可能的分组列（包含"区"、"地区"、"区域"等关键词）
        group_col_idx = None
        for i, header in enumerate(headers):
            if any(keyword in header for keyword in ["区", "地区", "区域", "类别", "分类"]):
                group_col_idx = i
                break

        if group_col_idx is None:
            return None

        # 按分组列的值分组
        groups = {}
        for row in table.rows[1:]:  # 跳过表头
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) <= group_col_idx:
                continue

            group_key = cells[group_col_idx]
            if not group_key:
                group_key = "其他"

            if group_key not in groups:
                groups[group_key] = []
            groups[group_key].append(cells)

        # 如果分组太少（<3个）或太多（>20个），不分组
        if len(groups) < 3 or len(groups) > 20:
            return None

        # 生成分组后的 Markdown
        markdown_lines = []
        for group_name, rows in sorted(groups.items()):
            markdown_lines.append(f"### {group_name}\n\n")

            # 表头
            markdown_lines.append("| " + " | ".join(headers) + " |\n")
            markdown_lines.append("| " + " | ".join(["---"] * len(headers)) + " |\n")

            # 数据行
            for cells in rows:
                escaped_cells = [cell.replace("|", "\\|") for cell in cells]
                markdown_lines.append("| " + " | ".join(escaped_cells) + " |\n")

            markdown_lines.append("\n")

        return markdown_lines
